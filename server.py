#!/usr/bin/env python3
import html
import io
import json
import os
import re
import shutil
import tempfile
import urllib.parse
import base64
import secrets
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path

import pdfplumber
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUTPUTS = ROOT / "public"
PORT = int(os.environ.get("PORT") or os.environ.get("RESUME_AGENT_PORT", "8766"))
APP_PASSWORD = os.environ.get("APP_PASSWORD", "")
APP_USERNAME = os.environ.get("APP_USERNAME", "admin")

LEXICON = [
    "AI", "Agent", "LLM", "RAG", "Prompt", "数据分析", "增长", "商业化", "B端",
    "SaaS", "产品策略", "用户研究", "A/B测试", "SQL", "Python", "机器学习",
    "深度学习", "模型评估", "推荐系统", "搜索", "支付", "风控", "CRM",
    "低代码", "自动化", "GTM", "运营", "战略", "跨团队", "项目管理",
    "0到1", "指标体系", "留存", "转化", "收入", "成本", "效率", "合规",
    "沟通", "组织协调", "抗压", "出差", "加班", "客户", "交付", "管理",
]

STOPWORDS = set(
    "岗位 职责 任职 要求 优先 以及 相关 负责 进行 具备 熟悉 能够 以上 工作 公司 团队 业务 产品 项目 经验 能力 包括 通过 对于 使用 一个 一份 我们 你将 需要".split()
)


def tokenize(text):
    found = []
    lower = text.lower()
    for word in LEXICON:
        if word.lower() in lower:
            found.append(word)
    words = re.findall(r"[A-Za-z][A-Za-z+#.]{1,28}|[\u4e00-\u9fa5]{2,8}", text)
    for word in words:
        if word not in STOPWORDS and len(word) > 1:
            found.append(word)
    counts = {}
    for word in found:
        counts[word] = counts.get(word, 0) + 1
    return [word for word, _ in sorted(counts.items(), key=lambda item: (-item[1], len(item[0]), item[0]))[:18]]


def normalize(text):
    return re.sub(r"\n{3,}", "\n\n", text.replace("\r", "").replace("\t", " ")).strip()


def extract_docx(path):
    doc = Document(path)
    lines = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def extract_pdf(path):
    lines = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                lines.append(text.strip())
    return "\n\n".join(lines)


def extract_text(path, filename):
    suffix = Path(filename).suffix.lower()
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".txt":
        return Path(path).read_text(encoding="utf-8", errors="ignore")
    raise ValueError("仅支持 .docx、.pdf、.txt。旧版 .doc 请先另存为 .docx。")


def score_line(line, keywords):
    lower = line.lower()
    score = sum(4 for kw in keywords if kw.lower() in lower)
    if re.search(r"[0-9%万亿kK+]", line):
        score += 3
    if re.search(r"负责|主导|搭建|推动|优化|提升|增长|降低|设计|上线|交付|管理|协作|沟通", line):
        score += 2
    return score


def intensity_count(intensity):
    return {
        "low": 3,
        "medium": 4,
        "high": 5,
        "conservative": 3,
        "balanced": 4,
        "aggressive": 5,
    }.get(intensity, 4)


def clean_company_type(company_type):
    return re.split(r"[（(]", company_type or "目标公司类型", maxsplit=1)[0].strip()


def split_resume_sections(resume_text):
    section_aliases = {
        "internship": re.compile(r"(实习经历|工作经历|工作/实习经历|实习实践|实践经历|职业经历|校外实践)"),
        "project": re.compile(r"(项目经历|项目经验|研究项目|课程项目|校园项目|调研项目|竞赛项目)"),
    }
    current = None
    sections = {"internship": [], "project": []}
    for raw in normalize(resume_text).splitlines():
        line = raw.strip()
        if not line:
            continue
        compact = re.sub(r"\s+", "", line)
        matched_heading = False
        for name, pattern in section_aliases.items():
            if len(compact) <= 12 and pattern.search(compact):
                current = name
                matched_heading = True
                break
        if matched_heading:
            continue
        if current:
            sections[current].append(line.strip(" -*\u2022"))

    all_lines = [line.strip(" -*\u2022") for line in normalize(resume_text).splitlines() if len(line.strip()) > 8]
    if not sections["internship"]:
        sections["internship"] = [
            line for line in all_lines
            if re.search(r"实习|公司|财务|运营|商务|助理|银行|客户|审核|发票|单据|合同|报销|数据|对账|部门", line)
        ][:8]
    if not sections["project"]:
        sections["project"] = [
            line for line in all_lines
            if re.search(r"项目|调研|研究|竞品|分析|报告|校园|课程|活动|社团|访谈|用户|市场|战略", line)
        ][:8]
    return sections


def infer_capabilities(text):
    capability_rules = [
        ("细节意识与流程规范", r"发票|报关|单据|合同|审核|审批|报销|对账|资料|凭证|合规|风控"),
        ("数据整理与分析", r"Excel|SQL|Python|Stata|数据|报表|统计|指标|分析|VLOOKUP|透视表"),
        ("市场与竞品研究", r"市场|竞品|行业|用户|调研|访谈|问卷|战略|竞对|产品分析|行研"),
        ("项目推进与节点把控", r"项目|推进|节点|排期|执行|落地|统筹|进度|交付|复盘"),
        ("沟通表达与组织协调", r"沟通|协调|跨部门|跨团队|会议|对接|合作方|客户|供应商|政府|企业|团队"),
        ("文档撰写与资料沉淀", r"文档|材料|方案|提案|报告|纪要|PPT|Word|邮件|会议资料"),
        ("跨文化与英文沟通", r"英语|英文|IELTS|CET|海外|国际|跨文化|留学|全球|外贸"),
        ("游戏理解与玩家同理心", r"游戏|ACG|玩家|版本|玩法|社区|发行|商务|CP|用户体验|内容生态"),
    ]
    found = []
    for label, pattern in capability_rules:
        if re.search(pattern, text, re.I):
            found.append(label)
    return found


def cap_text(text, limit=76):
    text = re.sub(r"\s+", " ", text).strip(" ，。；;-")
    return text if len(text) <= limit else text[:limit - 1].rstrip("，；、 ") + "。"


def rewrite_experience_line(line, jd, requirements, keywords, section_kind):
    original = re.sub(r"^[-*\u2022]\s*", "", line).rstrip("。；;")
    all_text = "\n".join([original, jd, requirements])
    caps = infer_capabilities(all_text)
    if not caps:
        caps = (keywords or ["岗位匹配能力"])[:2]
    caps = caps[:2]
    if re.search(r"发票|报关|单据|合同|审核|审批|报销|对账|凭证", original):
        bridge = "细致度、流程规范意识和执行力"
    elif re.search(r"竞品|市场|调研|行业|战略|访谈|报告|用户", original):
        bridge = "信息搜集、竞品分析和结构化判断能力"
    elif re.search(r"Excel|SQL|Python|Stata|数据|报表|统计|指标", original, re.I):
        bridge = "数据敏感度、工具使用和问题拆解能力"
    elif re.search(r"沟通|协调|会议|对接|合作|团队|活动|社团|组织", original):
        bridge = "沟通协调、节点把控和落地推进能力"
    else:
        bridge = f"{'、'.join(caps)}等岗位相关能力"
    rewritten = f"{original}，体现{bridge}。"
    return cap_text(rewritten, 88)


def adapt_experience_sections(resume_text, jd, requirements, keywords, intensity):
    sections = split_resume_sections(resume_text)
    count = intensity_count(intensity)
    adapted = {}
    for kind in ("internship", "project"):
        lines = [line for line in sections[kind] if len(line.strip()) > 6]
        scored = sorted(((line, score_line(line, keywords)) for line in lines), key=lambda item: item[1], reverse=True)
        picked = [line for line, _ in scored[:count]] or lines[:count]
        adapted[kind] = [rewrite_experience_line(line, jd, requirements, keywords, kind) for line in picked]
    return adapted


def build_evaluation(resume_text, jd, requirements, role, company, company_type, keywords):
    all_text = resume_text + "\n" + jd + "\n" + requirements
    matched = [kw for kw in keywords if kw.lower() in resume_text.lower()]
    capabilities = infer_capabilities(all_text)
    skill_words = "、".join((matched or keywords)[:4]) or "岗位相关技能"
    capability_words = "、".join(capabilities[:4]) or "资料整理、沟通协调、项目执行"
    role_name = role or "目标岗位"
    company_name = company or "目标公司"
    company_type_name = clean_company_type(company_type)
    has_numbers = bool(re.search(r"[0-9%万亿kK+]", resume_text))
    quant = "具备一定量化成果意识" if has_numbers else "可继续补充量化成果"
    mobility = "可适应加班、出差及快节奏工作"
    if re.search(r"出差|加班|抗压|高压|吃苦|驻场|客户现场", all_text):
        mobility = "能够适应加班、出差及高强度协作节奏"
    if "游戏" in all_text or "游戏" in company_type_name:
        return (
            f"候选人长期关注游戏及ACG内容生态，具备玩家同理心和产品观察意识，系统学习过{skill_words}等相关知识，"
            f"能够从用户体验、市场反馈和项目执行角度理解{company_name}{role_name}工作。过往经历体现出{capability_words}等能力，"
            f"性格细致负责，表达清晰，组织协调和学习能力较强；抗压性较好，吃苦耐劳，{mobility}。"
        )
    return (
        f"候选人围绕{company_type_name}及{role_name}所需能力，系统学习过{skill_words}等相关知识，"
        f"掌握{capability_words}等技能，具备与{company_name}{role_name}匹配的执行经验。"
        f"其性格稳重细致，对相关业务有兴趣，表达清晰，组织协调和问题闭环意识较强；"
        f"同时{quant}，抗压性较好，吃苦耐劳，{mobility}。"
    )


def add_run(paragraph, text, bold=False, size=10.5, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def insert_before(paragraph, text="", style=None):
    new_p = paragraph.insert_paragraph_before(text)
    if style:
        new_p.style = style
    return new_p


def create_new_docx(out_path, evaluation, adapted_sections, role, company):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(42)
    section.bottom_margin = Pt(42)
    section.left_margin = Pt(50)
    section.right_margin = Pt(50)
    title = doc.add_paragraph()
    add_run(title, f"{company or '目标公司'}｜{role or '目标岗位'}｜适配版简历", bold=True, size=15, color=(23, 78, 166))
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    h = doc.add_paragraph()
    add_run(h, "综合评价", bold=True, size=12, color=(23, 78, 166))
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Pt(12)
    add_run(p, evaluation, size=10)
    for heading, key in (("实习经历", "internship"), ("项目经历", "project")):
        h = doc.add_paragraph()
        add_run(h, heading, bold=True, size=12, color=(23, 78, 166))
        items = adapted_sections.get(key) or ["原简历中未识别到明确内容，建议补充对应经历后重新生成。"]
        for item in items:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(12)
            add_run(p, item, size=9.5)
    doc.save(out_path)


def generate_docx(upload_path, filename, payload):
    source_text = extract_text(upload_path, filename)
    keywords = tokenize(payload["jd"] + "\n" + payload["requirements"])
    evaluation = build_evaluation(
        source_text,
        payload["jd"],
        payload["requirements"],
        payload["role"],
        payload["company"],
        payload["companyType"],
        keywords,
    )
    adapted_sections = adapt_experience_sections(
        source_text,
        payload["jd"],
        payload["requirements"],
        keywords,
        payload["intensity"],
    )
    safe_company = re.sub(r"[\\/:*?\"<>|]+", "-", payload["company"] or "target-company")
    safe_role = re.sub(r"[\\/:*?\"<>|]+", "-", payload["role"] or "target-role")
    out_path = Path(tempfile.gettempdir()) / f"{safe_company}-{safe_role}-resume.docx"
    create_new_docx(out_path, evaluation, adapted_sections, payload["role"], payload["company"])
    preview = "\n".join([
        "综合评价",
        evaluation,
        "",
        "实习经历",
        *[f"- {item}" for item in adapted_sections.get("internship", [])],
        "",
        "项目经历",
        *[f"- {item}" for item in adapted_sections.get("project", [])],
    ])
    return out_path, preview, keywords


class Handler(BaseHTTPRequestHandler):
    def is_authorized(self):
        if not APP_PASSWORD:
            return True
        auth = self.headers.get("Authorization", "")
        if not auth.startswith("Basic "):
            return False
        try:
            decoded = base64.b64decode(auth.removeprefix("Basic ").strip()).decode("utf-8")
            username, password = decoded.split(":", 1)
        except Exception:
            return False
        return secrets.compare_digest(username, APP_USERNAME) and secrets.compare_digest(password, APP_PASSWORD)

    def require_auth(self):
        if self.is_authorized():
            return True
        self.send_response(401)
        self.send_header("WWW-Authenticate", 'Basic realm="Resume Agent"')
        self.send_header("Content-Type", "text/plain; charset=utf-8")
        self.end_headers()
        self.wfile.write("需要输入访问账号和密码。".encode("utf-8"))
        return False

    def do_GET(self):
        if self.path == "/healthz":
            self.send_response(200)
            self.send_header("Content-Type", "text/plain; charset=utf-8")
            self.end_headers()
            self.wfile.write(b"ok")
            return
        if not self.require_auth():
            return
        parsed = urllib.parse.urlparse(self.path)
        path = parsed.path
        if path == "/":
            path = "/resume-agent.html"
        target = (OUTPUTS / path.lstrip("/")).resolve()
        if not str(target).startswith(str(OUTPUTS.resolve())) or not target.exists():
            self.send_error(404)
            return
        content_type = "text/html; charset=utf-8" if target.suffix == ".html" else "application/octet-stream"
        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.end_headers()
        self.wfile.write(target.read_bytes())

    def do_POST(self):
        if not self.require_auth():
            return
        if self.path != "/api/generate":
            self.send_error(404)
            return
        upload_path = None
        out_path = None
        try:
            import cgi

            form = cgi.FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={
                    "REQUEST_METHOD": "POST",
                    "CONTENT_TYPE": self.headers.get("Content-Type"),
                },
            )
            file_item = form["resumeFile"]
            filename = Path(file_item.filename or "resume").name
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as tmp:
                tmp.write(file_item.file.read())
                upload_path = tmp.name
            payload = {
                "company": form.getfirst("company", ""),
                "role": form.getfirst("role", ""),
                "companyType": form.getfirst("companyType", "AI 原生公司"),
                "intensity": form.getfirst("intensity", "balanced"),
                "jd": form.getfirst("jd", ""),
                "requirements": form.getfirst("requirements", ""),
            }
            out_path, preview, keywords = generate_docx(upload_path, filename, payload)
            data = out_path.read_bytes()
            encoded_preview = html.escape(preview)
            self.send_response(200)
            self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            self.send_header("Content-Disposition", f"attachment; filename={urllib.parse.quote(out_path.name)}")
            self.send_header("X-Preview", urllib.parse.quote(encoded_preview))
            self.send_header("X-Keywords", urllib.parse.quote(json.dumps(keywords, ensure_ascii=False)))
            self.send_header("Content-Length", str(len(data)))
            self.end_headers()
            self.wfile.write(data)
        except Exception as exc:
            message = str(exc).encode("utf-8")
            self.send_response(400)
            self.send_header("Content-Type", "text/plain; charset=utf-8")
            self.send_header("Content-Length", str(len(message)))
            self.end_headers()
            self.wfile.write(message)
        finally:
            for path in (upload_path, out_path):
                if path:
                    try:
                        Path(path).unlink(missing_ok=True)
                    except Exception:
                        pass


if __name__ == "__main__":
    OUTPUTS.mkdir(exist_ok=True)
    server = ThreadingHTTPServer(("", PORT), Handler)
    print(f"Resume Agent running on port {PORT}", flush=True)
    server.serve_forever()
